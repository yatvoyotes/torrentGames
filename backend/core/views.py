import docx
from django.contrib import messages
from django.shortcuts import HttpResponseRedirect
from django.shortcuts import render
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from docx import Document
from docx.shared import Pt

from backend.core.models import Games, Basket, Order, Categories, Review
from backend.core.forms import ReviewAddForm
from core.randomkey import randomkey
import subprocess
from backend.users.models import Profile
from django.db.models import Q


def index(request, category_id=None, filtered=None):
    query = request.GET.get('q')
    if category_id:
        category = Categories.objects.get(id=category_id)
        games = Games.objects.filter(categories=category)
    elif filtered == 1:
        games = Games.objects.filter(ram__lt=4)
    elif filtered == 2:
        games = Games.objects.filter(release_date__year='2022')
    elif filtered == 3:
        games = Games.objects.order_by('-discount')
    elif query:
        games = Games.objects.filter(game_name__icontains=query)
        if games:
            pass
        else:
            messages.warning(request, f'По вашему запросу не найдено ни одной игры')
    else:
        games = Games.objects.all()
    paginator = Paginator(games, 10)
    page = request.GET.get('page')
    page_number = page
    page_obj = paginator.get_page(page_number)
    try:
        games = paginator.page(page)
    except PageNotAnInteger:
        games = paginator.page(1)
    except EmptyPage:
        games = paginator.page(paginator.num_pages)

    return render(request, 'mainContent.html', {'title': 'Главная страница сайта',
                                                'games': games,
                                                'paginator': paginator,
                                                'page_obj': page_obj,
                                                })


def game_detail(request, pk):
    games = Games.objects.get(id=pk)
    reviews = Review.objects.filter(game=pk, parent=None)
    reviewsChs = Review.objects.filter(~Q(parent=None), game=pk)
    profile = Profile.objects.all()
    return render(request, 'game_detail.html', {'games': games,
                                                'reviews': reviews,
                                                'reviewsChs': reviewsChs,
                                                'profile': profile})


def add_review(request, pk):
    form = ReviewAddForm(request.POST)
    game = Games.objects.get(id=pk)
    if form.is_valid():
        form = form.save(commit=False)
        if request.POST.get("parent", None):
            form.parent_id = int(request.POST.get("parent"))
        form.game_id = game.id
        form.name_id = request.user.id
        form.save()
        messages.success(request, f'Ваш отзыв успешно добавлен')
    return HttpResponseRedirect(request.META['HTTP_REFERER'])


def basket_add(request, game_id):
    game = Games.objects.get(id=game_id)
    baskets = Basket.objects.filter(user_id=request.user.id, game=game)
    if not baskets.exists():
        Basket.objects.create(user_id=request.user.id, game_id=game_id, quantity=1)
        messages.success(request, f'Товар добавлен в корзину')
    else:
        basket = baskets.first()
        basket.quantity += 1
        basket.save()
        messages.success(request, f'Товар добавлен в корзину')
    return HttpResponseRedirect(request.META['HTTP_REFERER'])


def basket_remove(request, basket_id, game_id=None):
    basket = Basket.objects.get(id=basket_id)
    if game_id and basket.quantity > 1:
        basket.quantity -= 1
        basket.save()
        messages.success(request, f'Копия игры удалена')
    else:
        basket.delete()
        messages.success(request, f'Товар удален из вашей корзины')
    return HttpResponseRedirect(request.META['HTTP_REFERER'])


def order_add(request, total):
    baskets = Basket.objects.filter(user_id=request.user.id)
    games = Games.objects.all()
    data = {}
    dict_num = 1
    for basket in baskets:
        diction = {
            "user": basket.user.user.username,
            "game": basket.game.game_name,
            "quantity": basket.quantity,
        }
        for game in games:
            if game.id == basket.game_id:
                price = game.price_with_discount = int(game.price - ((game.price / 100) * game.discount))
                diction["price"] = price
                diction["total"] = total
        data[dict_num] = diction
        dict_num += 1

    Order.objects.create(initiator_id=request.user.id, basket_history=data)
    Basket.objects.filter(user_id=request.user.id).delete()
    return HttpResponseRedirect(request.META['HTTP_REFERER'])


def orderpdf(request, order_id):
    doc = docx.Document()
    orders = Order.objects.filter(id=order_id)
    # добавляем первый параграф
    doc.add_paragraph("     Заказ №" + str(order_id))

    orders = Order.objects.filter(id=order_id)
    for order in orders:
        for item in order.basket_history:
            for game in order.basket_history[item]:

                if game == "game":
                    doc.add_paragraph(order.basket_history[item][game])
                if game == "quantity":
                    for i in range(0, order.basket_history[item][game]):
                        doc.add_paragraph(randomkey())
                # if game == "price":
                #      doc.add_paragraph(str(order.basket_history[item][game]))
                # if game == "total":
                #     total = order.basket_history[item][game]

    # doc.add_paragraph(str(total))

    doc.save('out.docx')
    subprocess.Popen('out.docx', shell=True)

    return HttpResponseRedirect(request.META['HTTP_REFERER'])

    # /Настройка шрифтов и добавление страницы
    # pdf = FPDF()
    # pdf.add_page()
    # pdf.add_font('DejaVu', '',
    #              'C:/Users/yurga/PycharmProjects/torrentGames/frontend/static/font/DejaVuSansCondensed.ttf',
    #              uni=True)
    # pdf.set_font('DejaVu', '', 12)
    # col_width = pdf.w / 4.5
    # row_height = pdf.font_size
    # spacing = 3
    #
    # # /Шапка
    # pdf.cell(200, 10,
    #          txt="-  --  --  --  --  --  --  --  --  --  --  --  --  --  --  --  --  --  --  --  --  --  --  --  -",
    #          ln=1, align="C")
    # pdf.cell(200, 10, txt="Кассовый чек", ln=1, align="C")
    # pdf.cell(200, 7, txt="Заказ №" + str(order_id), ln=1, align="C")
    #
    # # Шапка таблицы
    # pdf.cell(col_width, row_height * spacing,
    #          txt="Название", border=1, align="C")
    # pdf.cell(col_width, row_height * spacing,
    #          txt="Количество", border=1, align="C")
    # pdf.cell(col_width, row_height * spacing,
    #          txt="Ключ", border=1, align="C")
    # pdf.cell(col_width, row_height * spacing,
    #          txt="Цена за ед.", border=1, align="C")
    # pdf.ln(row_height * spacing)
    #
    # # Тело таблицы
    # orders = Order.objects.filter(id=order_id)
    # for order in orders:
    #     for item in order.basket_history:
    #         for game in order.basket_history[item]:
    #
    #             if game == "game":
    #                 pdf.cell(col_width, row_height * spacing,
    #                          txt=order.basket_history[item][game], border=1, align="C")
    #
    #             if game == "quantity" and order.basket_history[item][game] == 1:
    #                 pdf.cell(col_width, row_height * spacing,
    #                          txt=str(order.basket_history[item][game]), border=1, align="C")
    #                 for i in range(0, order.basket_history[item][game]):
    #                     if order.basket_history[item][game] == 1:
    #                         pdf.cell(col_width, row_height * spacing,
    #                                  txt=randomkey(), border=1, align="C")
    #             if game == "quantity" and order.basket_history[item][game] != 1:
    #                 pdf.cell(col_width, row_height * spacing,
    #                          txt=str(order.basket_history[item][game]), border=1, align="C")
    #                 for i in range(0, order.basket_history[item][game]):
    #                     if i == 0:
    #                         pdf.cell(col_width, row_height * spacing,
    #                                  txt=randomkey(), border=1, align="C")
    #                     else:
    #                         pdf.cell(col_width, row_height * spacing,
    #                                  border=1)
    #                         pdf.ln(row_height * spacing)
    #                         pdf.cell(col_width, row_height * spacing,
    #                                  border=1)
    #                         pdf.cell(col_width, row_height * spacing,
    #                                  border=1)
    #                         pdf.cell(col_width, row_height * spacing,
    #                                  txt=randomkey(), border=1, align="C")
    #             if game == "price":
    #                 pdf.cell(col_width, row_height * spacing,
    #                          txt=str(order.basket_history[item][game]) + " р.", border=1, align="C")
    #             if game == "total":
    #                 total = order.basket_history[item][game]
    #         pdf.ln(row_height * spacing)
    #
    # pdf.cell(170, 10, txt="ИТОГО " + str(total), ln=1, align="R")
    # pdf.set_text_color(0, 110, 51)
    # pdf.cell(170, 10, txt="СТАТУС: ОПЛАЧЕНО", ln=1, align="R")
    #
    # # /Footer
    # pdf.set_text_color(0, 0, 0)
    # pdf.cell(200, 10,
    #          txt="-  --  --  --  --  --  --  --  --  --  --  --  --  --  --  --  --  --  --  --  --  --  --  --  -",
    #          ln=1, align="C")
    #
    # pdf.output("C:/Users/yurga/PycharmProjects/torrentGames/frontend/static/check.pdf")  # /Выгрузка
    # subprocess.Popen("C:/Users/yurga/PycharmProjects/torrentGames/frontend/static/check.pdf",
    #                  shell=True)  # /Открыть pdf
